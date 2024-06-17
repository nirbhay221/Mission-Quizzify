# pdf_processing.py

# Necessary imports
import streamlit as st
from langchain_community.document_loaders import PyPDFLoader,Docx2txtLoader,UnstructuredPowerPointLoader, TextLoader, UnstructuredExcelLoader, WebBaseLoader
import os
import tempfile
import sys
import uuid
import os
from docx import Document as DocumentFromDocx
from dotenv import load_dotenv
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain_community.document_loaders import YoutubeLoader,UnstructuredURLLoader,PlaywrightURLLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from youtube_transcript_api import YouTubeTranscriptApi
import pandas as pd
print(sys.path)

load_dotenv()

class Document:
    def __init__(self, page_content='',metadata = None):
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {}
        
    def __str__(self):
        return f"Document(page_content='{self.page_content}', metadata={self.metadata})"
        
class DocumentProcessor:
    
    """
    This class encapsulates the functionality for processing uploaded PDF documents using Streamlit
    and Langchain's PyPDFLoader. It provides a method to render a file uploader widget, process the
    uploaded PDF files, extract their pages, and display the total number of pages extracted.
    """
    def __init__(self):
        self.pages = []  # List to keep track of pages from all documents
    
    def ingest_documents(self,input_type):
        """
        Renders a file uploader in a Streamlit app, processes uploaded PDF files,
        extracts their pages, and updates the self.pages list with the total number of pages.
        
        Given:
        - Handling of temporary files with unique names to avoid conflicts.
        """
        
        # Step 1: Render a file uploader widget. 
        st.title("Document Processor")
        
        if input_type == "PDF":
            self.process_pdf()
        if input_type == "DOCX":
            self.process_docx()
        if input_type == "DOC":
            self.process_docx()
        if input_type == "PPT":
            self.process_pptx()
        if input_type == "PPTX":
            self.process_pptx()
        if input_type == "Google Sheets":
            self.process_googleSheets()
        if input_type == "VIDEO_URL":
            self.process_video_url()
        if input_type == "URL":
            self.process_url()
        if input_type == "CSV":
            self.process_csv()
        if input_type == "TXT":
            self.process_txt()
        if input_type == "Excel":
            self.process_xlsx()
                
    def process_pdf(self):

        uploaded_files = st.file_uploader(
            label = "Streamlit PDF Uploader",
            accept_multiple_files= True,
            type = ["pdf"]
        )
        
        if uploaded_files is not None:
            pages = []
            for uploaded_file in uploaded_files:
                # Generate a unique identifier to append to the file's original name
                unique_id = uuid.uuid4().hex
                original_name, file_extension = os.path.splitext(uploaded_file.name)
                temp_file_name = f"{original_name}_{unique_id}{file_extension}"
                temp_file_path = os.path.join(tempfile.gettempdir(), temp_file_name)

                # Write the uploaded PDF to a temporary file
                with open(temp_file_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())

                # Step 2: Process the temporary file
                #####################################
                loader = PyPDFLoader(temp_file_path)
                docs = loader.load()
                all_pages = {}
                selected_pages = st.multiselect(
                    f'Select pages to process from {original_name}:',
                    options = [i for i, doc in enumerate(docs)],
                    format_func = lambda x: f'Page {x+1}'
                )
                if selected_pages:
                    selected_pages = [docs[i] for i in selected_pages]
                    all_pages[original_name] = selected_pages
                
                if all_pages:
                    self.pages = []
                    for pdf_name, pages in all_pages.items():
                        print("---PAGES---",pages)
                        self.pages.extend(pages)
                else:
                    self.pages.extend(docs)
                # Clean up by deleting the temporary file.
                os.unlink(temp_file_path)
            
            # Display the total number of pages processed.
            st.write(f"Total pages processed: {len(self.pages)}")
            self.pages = pages
            
            
            
    def process_xlsx(self):

        uploaded_files = st.file_uploader(
            label = "Streamlit Excel Uploader",
            accept_multiple_files= True,
            type = ["xlsx"]
        )
        
        if uploaded_files is not None:
            pages = []
            for uploaded_file in uploaded_files:
                # Generate a unique identifier to append to the file's original name
                unique_id = uuid.uuid4().hex
                original_name, file_extension = os.path.splitext(uploaded_file.name)
                temp_file_name = f"{original_name}_{unique_id}{file_extension}"
                temp_file_path = os.path.join(tempfile.gettempdir(), temp_file_name)

                # Write the uploaded PDF to a temporary file
                with open(temp_file_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())

                # Step 2: Process the temporary file
                #####################################
                loader = UnstructuredExcelLoader(temp_file_path)
                docs = loader.load()
                print("---------------DOCS-------------",docs)
                # Step 3: Then, Add the extracted pages to the 'pages' list.
                #####################################
                pages.extend(docs)
                # Clean up by deleting the temporary file.
                os.unlink(temp_file_path)
            
            # Display the total number of pages processed.
            st.write(f"Total pages processed: {len(self.pages)}")
            self.pages = pages
                     
    def process_txt(self):

        uploaded_files = st.file_uploader(
            label = "Streamlit TXT Uploader",
            accept_multiple_files= True,
            type = ["txt"]
        )
        
        if uploaded_files is not None:
            pages = []
            for uploaded_file in uploaded_files:
                # Generate a unique identifier to append to the file's original name
                unique_id = uuid.uuid4().hex
                original_name, file_extension = os.path.splitext(uploaded_file.name)
                temp_file_name = f"{original_name}_{unique_id}{file_extension}"
                temp_file_path = os.path.join(tempfile.gettempdir(), temp_file_name)

                with open(temp_file_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())
                with open(temp_file_path,'r') as txt_file:
                    txt_content = txt_file.read()
                    words_per_page = 500
                    words = txt_content.split()
                    num_words = len(words)
                    num_pages = (num_words // words_per_page) + (1 if num_words % words_per_page > 0 else 0 )
                    print("NUM PAGES ----------->",num_pages)
                    for i in range(num_pages):
                        start_idx = i * words_per_page
                        end_idx = min((i+1)*words_per_page , num_words)
                        segment_content = ' '.join(words[start_idx:end_idx])
                        metadata = {'soure':'txt', 'file_name': uploaded_file.name, 'segment_number': i+1}
                        pages.append((segment_content,metadata))
                    
                # Clean up by deleting the temporary file.
                os.unlink(temp_file_path)
                print("---------pages----------",pages)
            # Display the total number of pages processed.
            st.write(f"Total pages processed: {len(self.pages)}")
            self.pages = pages
            

    def process_csv(self):
        uploaded_files = st.file_uploader(
            label="Streamlit CSV Uploader",
            accept_multiple_files=True,
            type=["csv"]
        )

        if uploaded_files is not None:
            pages = []
            for uploaded_file in uploaded_files:
                unique_id = uuid.uuid4().hex
                original_name, file_extension = os.path.splitext(uploaded_file.name)
                temp_file_name = f"{original_name}_{unique_id}{file_extension}"
                temp_file_path = os.path.join(tempfile.gettempdir(), temp_file_name)

                with open(temp_file_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())

                
                df = pd.read_csv(temp_file_path)
                st.write("Data Preview:", df.head())

                st.write("Select rows and columns to include:")

                selected_rows = st.multiselect(
                    "Select specific rows:",
                    df.index,
                    format_func=lambda x: f"Row {x}"
                )

                row_range = st.slider(
                    "Select range of rows:",
                    0, len(df) - 1, (0, len(df) - 1)
                )

                selected_columns = st.multiselect(
                    "Select specific columns:",
                    df.columns
                )

                col_range = st.slider(
                    "Select range of columns:",
                    0, len(df.columns) - 1, (0, len(df.columns) - 1)
                )

                row_start, row_end = row_range
                col_start, col_end = col_range

                if not selected_rows and not selected_columns:
                    filtered_df = df.iloc[row_start:row_end + 1, col_start:col_end + 1]
                elif selected_rows and not selected_columns:
                    filtered_df = df.iloc[selected_rows, col_start:col_end + 1]
                elif selected_columns and not selected_rows:
                    column_names = df.columns[col_start:col_end + 1]
                    filtered_df = df.iloc[row_start:row_end + 1, df.columns.get_indexer(selected_columns)]
                else:
                    column_names = df.columns[col_start:col_end + 1]
                    selected_rows = range(row_start, row_end + 1)
                    filtered_df = df.iloc[selected_rows, df.columns.get_indexer(selected_columns)]

                st.write("Filtered Data Preview:", filtered_df)

                for index, row in filtered_df.iterrows():
                    filtered_text = row.to_string()
                    filtered_doc = {
                        "page_content": filtered_text,
                        "metadata": {
                            "row": index,
                            "columns": list(filtered_df.columns)
                        }
                    }
                    pages.append(filtered_doc)

                print("------FILTERED DOCS-----", filtered_df)
                print("---------------PAGES -------------", pages)
                os.unlink(temp_file_path)

            st.write(f"Total pages processed: {len(pages)}")
            self.pages = pages



    def process_url(self):
            urls_input = st.text_area("Enter URLs separated by comma (,)")
            urls = [url.strip() for url in urls_input.split(",") if url.strip()]
            
            if urls is not None:
                pages = []
                for url in urls:
                    try:    
                        loader = WebBaseLoader(url)
                        loader.request_kwargs = {'verify': False}
                        
                        docs = loader.load()
                    except Exception as e : 
                        st.write(f"Web Base Loader failed for {url} with error {str(e)}")
                        try:
                            loader = UnstructuredURLLoader(urls = [url])
                            docs = loader.load()
                        except Exception as e:
                            st.write(f"Unstructured URL Loader failed for {url} with error: {str(e)}")
                            try:
                                loader = PlaywrightURLLoader(urls = [url])
                                docs = loader.load()
                            except Exception as e:
                                st.write(f"Unstructured URL Loader failed for {url} with error: {str(e)}")
                                

                    pages.extend(docs)
                
                st.write(f"Total pages processed: {len(self.pages)}")
                self.pages = pages



                
    def process_docx(self):
        uploaded_files = st.file_uploader(
            label="Streamlit DOCX Uploader",
            accept_multiple_files=True,
            type=["docx", "doc"]
        )
        
        if uploaded_files is not None:
            pages = []
            page_number = 1
            
            for uploaded_file in uploaded_files:
                unique_id = uuid.uuid4().hex
                original_name, file_extension = os.path.splitext(uploaded_file.name)
                temp_file_name = f"{original_name}_{unique_id}{file_extension}"
                temp_file_path = os.path.join(tempfile.gettempdir(), temp_file_name)

                with open(temp_file_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())

                doc = DocumentFromDocx(temp_file_path)
                doc_text = []
                for paragraph in doc.paragraphs:
                    doc_text.append(paragraph.text)
                
                current_page = []
                word_count = 0 
                words_per_page = 500
                for para in doc_text:
                    word_count += len(para.split())
                    current_page.append(para)
                    if word_count >= words_per_page:
                        page_content = " ".join(current_page)
                        pages.append((page_content, {'source': 'docx', 'page_number': page_number}))
                        current_page = []
                        word_count = 0
                        page_number += 1
                if current_page:
                    page_content = " ".join(current_page)
                    pages.append((page_content, {'source': 'docx', 'page_number': page_number}))
                    page_number += 1
                
                os.unlink(temp_file_path)
            
            st.write(f"Total pages processed: {len(pages)}")
            
            st.markdown("### Select Pages")
            min_page = 1
            max_page = len(pages)
            page_range = st.slider(
                "Select pages to keep",
                min_value=min_page, max_value=max_page, value=(min_page, max_page)
            )
            
            filtered_pages = []
            for page_content, metadata in pages:
                if metadata['page_number'] in range(page_range[0], page_range[1] + 1):
                    filtered_pages.append((page_content, metadata))
            
            self.pages = filtered_pages
            
    def process_pptx(self):

        uploaded_files = st.file_uploader(
            label = "Streamlit PPTX Uploader",
            accept_multiple_files= True,
            type = ["pptx","ppt"]
        )
        
        if uploaded_files is not None:
            pages = []
            for uploaded_file in uploaded_files:
                # Generate a unique identifier to append to the file's original name
                unique_id = uuid.uuid4().hex
                original_name, file_extension = os.path.splitext(uploaded_file.name)
                temp_file_name = f"{original_name}_{unique_id}{file_extension}"
                temp_file_path = os.path.join(tempfile.gettempdir(), temp_file_name)

                # Write the uploaded PDF to a temporary file
                with open(temp_file_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())

                # Step 2: Process the temporary file
                #####################################
                loader = UnstructuredPowerPointLoader(temp_file_path)
                docs = loader.load()
                print("---------------DOCS-------------",docs)
                # Step 3: Then, Add the extracted pages to the 'pages' list.
                #####################################
                pages.extend(docs)
                # Clean up by deleting the temporary file.
                os.unlink(temp_file_path)
            
            # Display the total number of pages processed.
            st.write(f"Total pages processed: {len(self.pages)}")
            self.pages = pages



    def process_video_url(self):
        video_url = st.text_input("Enter Video URL", "")
        
        
        if video_url:
            
            video_id = video_url.split("v=")[1].split("&")[0]
            print(f'----------------LINK ID------------{video_id}')
            metadata = {'video_id': video_id}
            
            transcript = YouTubeTranscriptApi.get_transcript(video_id)
            timestamps = [entry['start'] for entry in transcript]
            selected_timestamp = st.slider("Select Timestamp", min_value=0, max_value=len(timestamps)-1)
            selected_entry = transcript[selected_timestamp]
            start_time = selected_entry['start']
            end_time = start_time + selected_entry['duration']
            transcript_content = ""
            docs = []
            
            pages = []
            i = 0
            for entry in transcript:
                if entry['start'] >= start_time and entry['start'] <= end_time:
                    words = entry['text'].split()
                    formatted_text = "\\n".join(words) 
                    metadata = {
                        'page' : i
                    }
                    transcript_content += formatted_text + " "
                    docs = [Document(page_content=transcript_content,metadata=metadata)]
                    
                    i+=1
                    pages.extend(docs)
            print(pages)
            
            # print(video_url)
            # loader = YoutubeLoader.from_youtube_url(video_url,add_video_info=True)
            # docs = loader.load()
            # print("-----------------------DOCS---------------------",docs)
            # # Add extracted pages to the 'pages' list
            
            # Display the total number of pages processed
            st.write(f"Total pages processed: {len(pages)}")
            self.pages = pages




if __name__ == "__main__":
    processor = DocumentProcessor()
    processor.ingest_documents("PDF")
